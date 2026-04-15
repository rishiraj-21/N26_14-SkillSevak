"""
Resume Parser 

Extracts text from PDF and DOCX resume files.

Features:
- PDF text extraction via pdfplumber
- DOCX text extraction via python-docx
- Text cleaning and normalization
- Section detection (skills, experience, education, etc.)
- Contact info extraction

Per PRD.md: Max file size 5MB, supports PDF/DOCX only.
"""

import re
import logging
from pathlib import Path
from typing import Dict, Optional, List

logger = logging.getLogger(__name__)


class ResumeParser:
    """
    Extract and process text from PDF/DOCX resumes.

    Usage:
        parser = ResumeParser()
        text = parser.extract_text('/path/to/resume.pdf')
        cleaned = parser.clean_text(text)
        sections = parser.detect_sections(text)

    Per PRD.md Section 7 (Phase 2):
    - Supports PDF and DOCX formats
    - Validates file size (max 5MB)
    - Extracts structured sections
    """

    # Section header patterns for detection
    # These are PATTERNS for detection, NOT a hardcoded skill list
    SECTION_PATTERNS = {
        'summary': r'(?i)^\s*(summary|objective|profile|about\s*me|professional\s*summary|career\s*objective)\s*:?\s*$',
        'skills': r'(?i)^\s*(skills|technical\s*skills|core\s*competencies|technologies|expertise|key\s*skills|areas\s*of\s*expertise)\s*:?\s*$',
        'experience': r'(?i)^\s*(experience|work\s*experience|work\s*history|employment|employment\s*history|professional\s*experience)\s*:?\s*$',
        'education': r'(?i)^\s*(education|academic|qualifications?|degrees?|academic\s*background)\s*:?\s*$',
        'projects': r'(?i)^\s*(projects|personal\s*projects|key\s*projects|portfolio)\s*:?\s*$',
        'certifications': r'(?i)^\s*(certifications?|certificates?|licenses?|credentials|professional\s*certifications?)\s*:?\s*$',
        'awards': r'(?i)^\s*(awards?|honors?|achievements?|accomplishments?|recognition)\s*:?\s*$',
        'languages': r'(?i)^\s*(languages?|language\s*proficiency)\s*:?\s*$',
        'interests': r'(?i)^\s*(interests?|hobbies|activities)\s*:?\s*$',
        'references': r'(?i)^\s*(references?|referees?)\s*:?\s*$',
    }

    # Max file size per PRD.md
    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB

    def __init__(self):
        """Initialize the parser."""
        self._pdfplumber = None
        self._docx = None

    def extract_text(self, file_path: str) -> str:
        """
        Main entry point for text extraction.

        Args:
            file_path: Path to PDF or DOCX file

        Returns:
            Extracted text content

        Raises:
            ValueError: If file type is unsupported or file too large
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        # Validate file size
        file_size = path.stat().st_size
        if file_size > self.MAX_FILE_SIZE:
            raise ValueError(
                f"File size ({file_size / 1024 / 1024:.1f}MB) exceeds "
                f"maximum allowed ({self.MAX_FILE_SIZE / 1024 / 1024}MB)"
            )

        ext = path.suffix.lower()

        if ext == '.pdf':
            return self._extract_pdf(str(path))
        elif ext == '.docx':
            return self._extract_docx(str(path))
        else:
            raise ValueError(
                f"Unsupported file type: {ext}. Only PDF and DOCX are supported."
            )

    def _extract_pdf(self, path: str) -> str:
        """
        Extract text from PDF using pdfplumber.

        Handles multi-page documents and tables.
        """
        try:
            import pdfplumber

            text_parts = []

            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages):
                    try:
                        # Extract main text
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)

                        # Also extract tables (common in resumes)
                        tables = page.extract_tables()
                        for table in tables:
                            for row in table:
                                if row:
                                    row_text = ' | '.join(
                                        cell.strip() if cell else ''
                                        for cell in row
                                    )
                                    if row_text.strip():
                                        text_parts.append(row_text)

                    except Exception as e:
                        logger.warning(f"Error extracting page {i + 1}: {e}")
                        continue

            extracted_text = '\n'.join(text_parts)
            logger.info(f"Extracted {len(extracted_text)} chars from PDF")
            return extracted_text

        except ImportError:
            logger.error("pdfplumber not installed. Run: pip install pdfplumber")
            raise
        except Exception as e:
            logger.error(f"PDF extraction failed for {path}: {e}")
            raise

    def _extract_docx(self, path: str) -> str:
        """
        Extract text from DOCX using python-docx.

        Handles paragraphs, tables, and headers.
        """
        try:
            from docx import Document

            doc = Document(path)
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)

            # Extract tables (skills often in tables)
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        text_parts.append(' | '.join(row_texts))

            extracted_text = '\n'.join(text_parts)
            logger.info(f"Extracted {len(extracted_text)} chars from DOCX")
            return extracted_text

        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"DOCX extraction failed for {path}: {e}")
            raise

    def clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.

        Operations:
        - Remove excessive whitespace
        - Normalize line breaks
        - Remove non-printable characters
        - Preserve structure for section detection
        """
        if not text:
            return ''

        # Replace tabs with spaces
        text = text.replace('\t', ' ')

        # Replace multiple spaces with single space (but keep newlines)
        text = re.sub(r'[ ]+', ' ', text)

        # Normalize multiple newlines to double newline (paragraph break)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)

        # Remove non-printable characters (keep basic punctuation and newlines)
        text = re.sub(r'[^\x20-\x7E\n\r\t]', '', text)

        # Remove empty lines at start/end
        text = text.strip()

        return text

    def detect_sections(self, text: str) -> Dict[str, str]:
        """
        Detect and extract resume sections.

        Returns dict with keys: contact, summary, skills, experience,
        education, projects, certifications, awards, other

        Per PRD.md: Section detection is pattern-based, but skill
        extraction within sections is NLP-based (Phase 3).
        """
        sections = {
            'contact': '',
            'summary': '',
            'skills': '',
            'experience': '',
            'education': '',
            'projects': '',
            'certifications': '',
            'awards': '',
            'languages': '',
            'interests': '',
            'references': '',
            'other': '',
        }

        if not text:
            return sections

        lines = text.split('\n')
        current_section = 'contact'  # First few lines usually contact info
        section_content: List[str] = []
        contact_line_count = 0
        max_contact_lines = 8  # Assume first 8 lines max are contact

        for line in lines:
            line_stripped = line.strip()

            # Skip empty lines but preserve them for content
            if not line_stripped:
                section_content.append('')
                continue

            # Check if this line is a section header
            new_section = self._identify_section_header(line_stripped)

            if new_section:
                # Save current section content
                content = '\n'.join(section_content).strip()
                if content:
                    sections[current_section] = content

                # Start new section
                section_content = []
                current_section = new_section

            else:
                # Add to current section
                section_content.append(line)

                # Track contact lines
                if current_section == 'contact':
                    contact_line_count += 1
                    # After max contact lines, switch to 'other' until we find a header
                    if contact_line_count >= max_contact_lines:
                        content = '\n'.join(section_content).strip()
                        if content:
                            sections['contact'] = content
                        section_content = []
                        current_section = 'other'

        # Save last section
        content = '\n'.join(section_content).strip()
        if content:
            sections[current_section] = content

        # Log detected sections
        detected = [k for k, v in sections.items() if v.strip()]
        logger.info(f"Detected sections: {detected}")

        return sections

    def _identify_section_header(self, line: str) -> Optional[str]:
        """
        Check if a line is a section header.

        Returns section name if match found, None otherwise.
        """
        if not line:
            return None

        # Headers are usually short (< 50 chars) and may be uppercase
        if len(line) > 60:
            return None

        # Check against patterns
        for section, pattern in self.SECTION_PATTERNS.items():
            if re.search(pattern, line):
                return section

        # Also check for ALL CAPS headers (common format)
        if line.isupper() and len(line) < 30:
            line_lower = line.lower()
            for section, pattern in self.SECTION_PATTERNS.items():
                # Check if the uppercase line matches any keyword
                keywords = ['summary', 'objective', 'skills', 'experience',
                           'education', 'projects', 'certifications', 'awards']
                for keyword in keywords:
                    if keyword in line_lower:
                        return section

        return None

    def extract_contact_info(self, text: str) -> Dict[str, str]:
        """
        Extract contact information from resume.

        Returns dict with: email, phone, linkedin, github, website
        """
        contact = {
            'email': '',
            'phone': '',
            'linkedin': '',
            'github': '',
            'website': '',
            'location': '',
        }

        if not text:
            return contact

        # Email pattern
        email_match = re.search(
            r'[\w.+-]+@[\w-]+\.[\w.-]+',
            text
        )
        if email_match:
            contact['email'] = email_match.group()

        # Phone pattern (various formats)
        phone_patterns = [
            r'\+?1?[-.\s]?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}',  # US
            r'\+?[0-9]{1,3}[-.\s]?[0-9]{4,5}[-.\s]?[0-9]{4,6}',  # International
        ]
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                contact['phone'] = phone_match.group().strip()
                break

        # LinkedIn URL
        linkedin_match = re.search(
            r'(?:https?://)?(?:www\.)?linkedin\.com/in/([\w-]+)',
            text,
            re.IGNORECASE
        )
        if linkedin_match:
            contact['linkedin'] = f"linkedin.com/in/{linkedin_match.group(1)}"

        # GitHub URL
        github_match = re.search(
            r'(?:https?://)?(?:www\.)?github\.com/([\w-]+)',
            text,
            re.IGNORECASE
        )
        if github_match:
            contact['github'] = f"github.com/{github_match.group(1)}"

        # Generic website (excluding social media)
        website_match = re.search(
            r'(?:https?://)?(?:www\.)?([\w-]+\.[\w.-]+)',
            text
        )
        if website_match and not any(x in website_match.group().lower()
                                     for x in ['linkedin', 'github', 'facebook', 'twitter']):
            contact['website'] = website_match.group()

        return contact

    def get_word_count(self, text: str) -> int:
        """Get word count for profile completeness calculation."""
        if not text:
            return 0
        return len(text.split())

    def get_section_stats(self, sections: Dict[str, str]) -> Dict[str, int]:
        """
        Get statistics about detected sections.

        Useful for profile completeness scoring.
        """
        stats = {}
        for section, content in sections.items():
            stats[f'{section}_words'] = self.get_word_count(content)
            stats[f'{section}_chars'] = len(content)
            stats[f'{section}_lines'] = len(content.split('\n')) if content else 0

        # Count non-empty sections
        stats['sections_filled'] = sum(1 for v in sections.values() if v.strip())
        stats['total_sections'] = len(sections)

        return stats


    def calculate_completeness_score(self, sections: Dict[str, str]) -> int:
        """
        Calculate resume completeness score (0-100).

        New Balanced Scoring:
        - Contact info: 10 points
        - Skills section: 25 points
        - Experience section: 20 points
        - Education section: 15 points
        - Summary/Objective: 5 points
        - Projects: 20 points
        - Other sections: 5 points
        """
        score = 0

        # Contact (10 points)
        if sections.get('contact', '').strip():
            contact_text = sections['contact']
            if re.search(r'@', contact_text):
                score += 5
            if re.search(r'\d{3}.*\d{3}.*\d{4}', contact_text):
                score += 5

        # Skills (25 points)
        skills_text = sections.get('skills', '')
        if skills_text.strip():
            word_count = self.get_word_count(skills_text)
            if word_count >= 20:
                score += 25
            elif word_count >= 6:
                score += 18
            elif word_count >= 3:
                score += 12
            else:
                score += 6

        # Experience (20 points)
        exp_text = sections.get('experience', '')
        if exp_text.strip():
            word_count = self.get_word_count(exp_text)
            if word_count >= 100:
                score += 20
            elif word_count >= 50:
                score += 15
            elif word_count >= 20:
                score += 10
            else:
                score += 5

        # Education (15 points)
        edu_text = sections.get('education', '')
        if edu_text.strip():
            word_count = self.get_word_count(edu_text)
            if word_count >= 20:
                score += 15
            elif word_count >= 10:
                score += 10
            else:
                score += 5

        # Summary (5 points)
        if sections.get('summary', '').strip():
            score += 5

        # Projects (20 points)
        proj_text = sections.get('projects', '')
        if proj_text.strip():
            word_count = self.get_word_count(proj_text)
            if word_count >= 40:
                score += 20
            elif word_count >= 20:
                score += 15
            else:
                score += 10

        # Other sections (5 points max)
        other_sections = ['certifications', 'awards', 'languages']
        for section in other_sections:
            if sections.get(section, '').strip():
                score += 2
                if score > 100:
                    score = 100
                    break

        return min(score, 100)
  

    # def calculate_completeness_score(self, sections: Dict[str, str]) -> int:
    #     """
    #     Calculate resume completeness score (0-100).

    #     Scoring:
    #     - Contact info present: 15 points
    #     - Skills section: 20 points
    #     - Experience section: 25 points
    #     - Education section: 15 points
    #     - Summary/Objective: 10 points
    #     - Projects: 10 points
    #     - Other sections: 5 points
    #     """
    #     score = 0

    #     # Contact (15 points)
    #     if sections.get('contact', '').strip():
    #         contact_text = sections['contact']
    #         # Check for email
    #         if re.search(r'@', contact_text):
    #             score += 10
    #         # Check for phone
    #         if re.search(r'\d{3}.*\d{3}.*\d{4}', contact_text):
    #             score += 5

    #     # Skills (20 points)
    #     skills_text = sections.get('skills', '')
    #     if skills_text.strip():
    #         word_count = self.get_word_count(skills_text)
    #         if word_count >= 20:
    #             score += 20
    #         elif word_count >= 10:
    #             score += 15
    #         elif word_count >= 5:
    #             score += 10
    #         else:
    #             score += 5

    #     # Experience (25 points)
    #     exp_text = sections.get('experience', '')
    #     if exp_text.strip():
    #         word_count = self.get_word_count(exp_text)
    #         if word_count >= 100:
    #             score += 25
    #         elif word_count >= 50:
    #             score += 20
    #         elif word_count >= 20:
    #             score += 15
    #         else:
    #             score += 10

    #     # Education (15 points)
    #     edu_text = sections.get('education', '')
    #     if edu_text.strip():
    #         word_count = self.get_word_count(edu_text)
    #         if word_count >= 20:
    #             score += 15
    #         elif word_count >= 10:
    #             score += 10
    #         else:
    #             score += 5

    #     # Summary (10 points)
    #     if sections.get('summary', '').strip():
    #         score += 10

    #     # Projects (10 points)
    #     if sections.get('projects', '').strip():
    #         score += 10

    #     # Other sections (5 points max)
    #     other_sections = ['certifications', 'awards', 'languages']
    #     for section in other_sections:
    #         if sections.get(section, '').strip():
    #             score += 2
    #             if score > 100:
    #                 score = 100
    #                 break

    #     return min(score, 100)
